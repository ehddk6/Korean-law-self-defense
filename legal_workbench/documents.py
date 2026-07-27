from __future__ import annotations

import html
import importlib
import mimetypes
import os
import re
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import asdict, dataclass
from datetime import date
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET


class DocumentError(RuntimeError):
    pass


@dataclass(slots=True)
class ExtractionResult:
    text: str
    media_type: str
    status: str
    confidence: float | None
    warnings: list[str]
    page_count: int | None = None

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm", ".eml"}


def extract_document(path: Path) -> ExtractionResult:
    source = Path(path)
    if not source.is_file():
        raise DocumentError(f"문서가 존재하지 않습니다: {source}")
    suffix = source.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        text, encoding = _read_text(source)
        return ExtractionResult(
            text=text,
            media_type=mimetypes.guess_type(source.name)[0] or "text/plain",
            status="extracted",
            confidence=1.0,
            warnings=[f"decoded-as:{encoding}"],
        )
    if suffix == ".docx":
        text = _extract_zip_xml(source, prefixes=("word/document.xml", "word/footnotes.xml", "word/endnotes.xml"))
        return ExtractionResult(text, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "extracted", 1.0, [])
    if suffix == ".hwpx":
        text = _extract_zip_xml(source, prefixes=("Contents/section",))
        return ExtractionResult(text, "application/vnd.hancom.hwpx", "extracted", 1.0, [])
    if suffix == ".pdf":
        return _extract_pdf(source)
    if suffix == ".hwp":
        raise DocumentError(
            "바이너리 HWP는 먼저 $hwpx Skill의 검증된 HWP→HWPX 변환을 수행해야 합니다. "
            "변환된 HWPX만 ingest 하십시오."
        )
    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        return ExtractionResult(
            text="",
            media_type=mimetypes.guess_type(source.name)[0] or "application/octet-stream",
            status="needs_ocr",
            confidence=0.0,
            warnings=["이미지 OCR이 필요하며 OCR 검증 전에는 증거 인용에 사용할 수 없습니다."],
        )
    raise DocumentError(f"지원하지 않는 문서 형식입니다: {suffix or '<no extension>'}")


def _read_text(path: Path) -> tuple[str, str]:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-16", "cp949"):
        try:
            return data.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace"), "utf-8-replace"


def _extract_zip_xml(path: Path, *, prefixes: tuple[str, ...]) -> str:
    if not zipfile.is_zipfile(path):
        raise DocumentError(f"유효한 ZIP 기반 문서가 아닙니다: {path.name}")
    paragraphs: list[str] = []
    with zipfile.ZipFile(path) as archive:
        names = [name for name in archive.namelist() if any(name.startswith(prefix) for prefix in prefixes)]
        for name in sorted(names):
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError as exc:
                raise DocumentError(f"XML 파싱 실패: {name}: {exc}") from exc
            current: list[str] = []
            for element in root.iter():
                tag = element.tag.rsplit("}", 1)[-1]
                if tag in {"t", "delText", "instrText"} and element.text:
                    current.append(element.text)
                elif tag in {"p", "tr"} and current:
                    line = "".join(current).strip()
                    if line:
                        paragraphs.append(line)
                    current = []
            if current:
                line = "".join(current).strip()
                if line:
                    paragraphs.append(line)
    return "\n".join(paragraphs)


def _extract_pdf(path: Path) -> ExtractionResult:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentError("PDF 추출에는 pypdf가 필요합니다.") from exc
    reader = PdfReader(str(path))
    pages: list[str] = []
    empty_pages = 0
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            empty_pages += 1
        pages.append(f"\n--- page {index} ---\n{text.strip()}")
    page_count = len(reader.pages)
    if page_count == 0 or empty_pages == page_count:
        return ExtractionResult(
            text="",
            media_type="application/pdf",
            status="needs_ocr",
            confidence=0.0,
            warnings=["텍스트가 없는 스캔 PDF입니다. OCR 및 원문 화면 대조가 필요합니다."],
            page_count=page_count,
        )
    confidence = max(0.0, 1.0 - empty_pages / max(page_count, 1))
    warnings = []
    if empty_pages:
        warnings.append(f"텍스트가 추출되지 않은 페이지: {empty_pages}/{page_count}")
    return ExtractionResult(
        text="\n".join(pages).strip(),
        media_type="application/pdf",
        status="extracted" if not empty_pages else "partial",
        confidence=confidence,
        warnings=warnings,
        page_count=page_count,
    )


def markdown_blocks(markdown: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line:
            blocks.append(("blank", ""))
        elif line.startswith("### "):
            blocks.append(("heading3", line[4:]))
        elif line.startswith("## "):
            blocks.append(("heading2", line[3:]))
        elif line.startswith("# "):
            blocks.append(("heading1", line[2:]))
        elif re.match(r"^\s*[-*]\s+", line):
            blocks.append(("bullet", re.sub(r"^\s*[-*]\s+", "", line)))
        elif re.match(r"^\s*\d+[.)]\s+", line):
            blocks.append(("number", re.sub(r"^\s*\d+[.)]\s+", "", line)))
        elif line.startswith("> "):
            blocks.append(("quote", line[2:]))
        else:
            blocks.append(("paragraph", line))
    return blocks


def presentation_blocks(markdown: str, title: str) -> list[tuple[str, str]]:
    """Remove a leading Markdown title already rendered by the output format."""
    blocks = markdown_blocks(markdown)
    first_content = next((index for index, (kind, _) in enumerate(blocks) if kind != "blank"), None)
    if first_content is None:
        return blocks
    kind, text = blocks[first_content]
    if kind == "heading1" and text.strip().casefold() == title.strip().casefold():
        del blocks[first_content]
        if first_content < len(blocks) and blocks[first_content][0] == "blank":
            del blocks[first_content]
    return blocks


def create_docx(markdown: str, output: Path, *, title: str) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt
    except ImportError as exc:
        raise DocumentError("DOCX 생성에는 python-docx가 필요합니다.") from exc
    destination = Path(output)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    styles = document.styles
    for style_name, size, bold in (("Normal", 10, False), ("Title", 18, True), ("Heading 1", 15, True), ("Heading 2", 13, True), ("Heading 3", 11, True)):
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = bold
    title_paragraph = document.add_paragraph(title, style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning = document.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = warning.add_run("사용자 검토용 초안 - 법원 제출 완료 문서가 아님")
    run.bold = True
    for kind, text in presentation_blocks(markdown, title):
        if kind == "blank":
            document.add_paragraph("")
        elif kind.startswith("heading"):
            level = int(kind[-1])
            document.add_heading(text, level=level)
        elif kind == "bullet":
            document.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            document.add_paragraph(text, style="List Number")
        elif kind == "quote":
            paragraph = document.add_paragraph(text)
            paragraph.paragraph_format.left_indent = Mm(8)
        else:
            paragraph = document.add_paragraph(text)
            paragraph.paragraph_format.line_spacing = 1.6
    footer = section.footer.paragraphs[0]
    footer.text = "한국법 자가소송 워크벤치 | 사용자 최종 확인 필요"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(destination)
    validate_docx(destination)
    return destination


def create_pdf(markdown: str, output: Path, *, title: str) -> Path:
    try:
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise DocumentError("PDF 생성에는 reportlab이 필요합니다.") from exc
    destination = Path(output)
    destination.parent.mkdir(parents=True, exist_ok=True)
    regular = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "malgun.ttf"
    bold = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "malgunbd.ttf"
    font_name = "Helvetica"
    bold_name = "Helvetica-Bold"
    if regular.exists():
        pdfmetrics.registerFont(TTFont("MalgunGothic", str(regular)))
        font_name = "MalgunGothic"
    if bold.exists():
        pdfmetrics.registerFont(TTFont("MalgunGothicBold", str(bold)))
        bold_name = "MalgunGothicBold"
    styles = getSampleStyleSheet()
    body = ParagraphStyle("LegalBody", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=15, spaceAfter=4)
    heading = {
        1: ParagraphStyle("LegalH1", parent=body, fontName=bold_name, fontSize=15, leading=21, spaceBefore=9, spaceAfter=5),
        2: ParagraphStyle("LegalH2", parent=body, fontName=bold_name, fontSize=13, leading=18, spaceBefore=8, spaceAfter=5),
        3: ParagraphStyle("LegalH3", parent=body, fontName=bold_name, fontSize=11, leading=16, spaceBefore=7, spaceAfter=4),
    }
    title_style = ParagraphStyle("LegalTitle", parent=body, fontName=bold_name, fontSize=18, leading=25, alignment=TA_CENTER, spaceAfter=10)
    warning_style = ParagraphStyle("LegalWarning", parent=body, fontName=bold_name, fontSize=9, leading=13, alignment=TA_CENTER, spaceAfter=12)
    story = [
        Paragraph(html.escape(title), title_style),
        Paragraph("사용자 검토용 초안 - 법원 제출 완료 문서가 아님", warning_style),
    ]
    for kind, text in presentation_blocks(markdown, title):
        escaped = html.escape(text)
        if kind == "blank":
            story.append(Spacer(1, 2 * mm))
        elif kind.startswith("heading"):
            story.append(Paragraph(escaped, heading[int(kind[-1])]))
        elif kind == "bullet":
            story.append(Paragraph(f"- {escaped}", body))
        elif kind == "number":
            story.append(Paragraph(escaped, body))
        elif kind == "quote":
            quote_style = ParagraphStyle("LegalQuote", parent=body, leftIndent=8 * mm, textColor="#444444")
            story.append(Paragraph(escaped, quote_style))
        else:
            story.append(Paragraph(escaped, body))

    def footer(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.drawCentredString(A4[0] / 2, 10 * mm, f"사용자 최종 확인 필요 | {doc.page}")
        canvas.restoreState()

    pdf = SimpleDocTemplate(
        str(destination),
        pagesize=A4,
        leftMargin=25 * mm,
        rightMargin=25 * mm,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
        title=title,
        author="Korean Legal Workbench",
    )
    pdf.build(story, onFirstPage=footer, onLaterPages=footer)
    validate_pdf(destination)
    return destination


def create_hwpx(markdown: str, output: Path, *, title: str) -> Path:
    skill_dir = Path(
        os.environ.get("HWPX_SKILL_DIR", str(Path.home() / ".codex" / "skills" / "hwpx-skill"))
    ).resolve()
    helpers_dir = skill_dir / "scripts"
    if not (helpers_dir / "hwpx_helpers.py").is_file():
        raise DocumentError(f"HWPX Skill을 찾을 수 없습니다: {skill_dir}")
    sys.path.insert(0, str(helpers_dir))
    try:
        helpers = importlib.import_module("hwpx_helpers")
        helpers.reset_id()
        header = skill_dir / "templates" / "government" / "header.xml"
        reference = skill_dir / "assets" / "government-reference.hwpx"
        helpers.validate_header_for_government(header)
        secpr, colpr = helpers.extract_secpr_and_colpr(reference)
        parts = [
            '<?xml version="1.0" encoding="UTF-8" standalone="yes" ?>',
            f"<hs:sec {helpers.NS_DECL}>",
            helpers.make_first_para(secpr, colpr),
        ]
        parts.extend(helpers.make_cover_page(title, subtitle="사용자 검토용 초안", date=date.today().strftime("%Y. %m.")))
        section_number = 1
        for kind, text in presentation_blocks(markdown, title):
            if kind == "heading1" or kind == "heading2":
                parts.append(helpers.make_section_bar(str(section_number), text))
                section_number += 1
            elif kind == "blank":
                parts.append(helpers.make_empty_line())
            elif text:
                marker = "-" if kind in {"bullet", "number"} else ""
                parts.append(helpers.make_body_para(marker, text))
        parts.append("</hs:sec>")
        destination = Path(output)
        destination.parent.mkdir(parents=True, exist_ok=True)
        with tempfile.TemporaryDirectory(prefix="legal-hwpx-") as temp_dir:
            section_file = Path(temp_dir) / "section0.xml"
            section_file.write_text("\n".join(parts), encoding="utf-8")
            commands = [
                [
                    sys.executable,
                    str(helpers_dir / "build_hwpx.py"),
                    "--header",
                    str(header),
                    "--section",
                    str(section_file),
                    "--title",
                    title,
                    "--output",
                    str(destination),
                ],
                [sys.executable, str(helpers_dir / "fix_namespaces.py"), str(destination)],
                [sys.executable, str(helpers_dir / "validate.py"), str(destination)],
            ]
            for command in commands:
                subprocess.run(command, check=True, capture_output=True, text=True, encoding="utf-8", errors="replace")
        validate_hwpx(destination)
        return destination
    except subprocess.CalledProcessError as exc:
        raise DocumentError(f"HWPX 생성 또는 검증 실패: {exc.stderr or exc.stdout}") from exc
    finally:
        if sys.path and sys.path[0] == str(helpers_dir):
            sys.path.pop(0)


def validate_docx(path: Path) -> dict[str, Any]:
    required = {"[Content_Types].xml", "word/document.xml"}
    return _validate_zip(path, required=required, expected_mimetype=None)


def validate_hwpx(path: Path) -> dict[str, Any]:
    required = {"mimetype", "Contents/header.xml", "Contents/section0.xml", "Contents/content.hpf"}
    result = _validate_zip(path, required=required, expected_mimetype="application/hwp+zip")
    with zipfile.ZipFile(path) as archive:
        first = archive.infolist()[0]
        if first.filename != "mimetype" or first.compress_type != zipfile.ZIP_STORED:
            raise DocumentError("HWPX mimetype은 첫 엔트리이자 ZIP_STORED여야 합니다.")
    return result


def validate_pdf(path: Path) -> dict[str, Any]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentError("PDF 검증에는 pypdf가 필요합니다.") from exc
    reader = PdfReader(str(path))
    if not reader.pages:
        raise DocumentError("PDF에 페이지가 없습니다.")
    return {"valid": True, "page_count": len(reader.pages)}


def _validate_zip(path: Path, *, required: set[str], expected_mimetype: str | None) -> dict[str, Any]:
    if not zipfile.is_zipfile(path):
        raise DocumentError(f"ZIP 기반 문서가 손상되었습니다: {path}")
    with zipfile.ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise DocumentError(f"ZIP CRC 오류: {bad}")
        names = set(archive.namelist())
        missing = sorted(required - names)
        if missing:
            raise DocumentError(f"필수 문서 파트 누락: {', '.join(missing)}")
        xml_names = [name for name in names if name.endswith(".xml") or name.endswith(".hpf")]
        for name in xml_names:
            try:
                ET.fromstring(archive.read(name))
            except ET.ParseError as exc:
                raise DocumentError(f"XML 검증 실패: {name}: {exc}") from exc
        if expected_mimetype:
            actual = archive.read("mimetype").decode("utf-8").strip()
            if actual != expected_mimetype:
                raise DocumentError(f"HWPX mimetype 불일치: {actual}")
    return {"valid": True, "entries": len(names)}


def rehydrate_document(source: Path, destination: Path, mapping: dict[str, str]) -> Path:
    suffix = Path(source).suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        text, _ = _read_text(Path(source))
        from .security import rehydrate_text

        restored = rehydrate_text(text, mapping)
        Path(destination).parent.mkdir(parents=True, exist_ok=True)
        Path(destination).write_text(restored, encoding="utf-8", newline="\n")
        return Path(destination)
    if suffix in {".docx", ".hwpx"}:
        return _rehydrate_zip_xml(Path(source), Path(destination), mapping)
    if suffix == ".pdf":
        raise DocumentError("PDF 텍스트 직접 복원은 지원하지 않습니다. 복원된 원본에서 다시 PDF를 생성하십시오.")
    raise DocumentError(f"복원할 수 없는 형식입니다: {suffix}")


def _rehydrate_zip_xml(source: Path, destination: Path, mapping: dict[str, str]) -> Path:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(delete=False, suffix=destination.suffix, dir=destination.parent) as handle:
        temporary = Path(handle.name)
    try:
        with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(temporary, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.endswith(".xml") or item.filename.endswith(".hpf"):
                    try:
                        text = data.decode("utf-8")
                    except UnicodeDecodeError:
                        pass
                    else:
                        for original, token in sorted(mapping.items(), key=lambda pair: -len(pair[1])):
                            text = text.replace(token, html.escape(original, quote=True))
                        data = text.encode("utf-8")
                zout.writestr(item, data)
        os.replace(temporary, destination)
    finally:
        if temporary.exists():
            temporary.unlink()
    if destination.suffix.lower() == ".docx":
        validate_docx(destination)
    else:
        validate_hwpx(destination)
    return destination
